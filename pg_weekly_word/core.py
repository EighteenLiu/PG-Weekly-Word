from __future__ import annotations

import json
import os
import re
import shutil
import sys
import tempfile
import zipfile
from collections import Counter, defaultdict
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime, timedelta
from functools import lru_cache
from pathlib import Path
from typing import Any, Iterable
import xml.etree.ElementTree as ET

import openpyxl
import xlrd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docxtpl import DocxTemplate


PROJECT_ROOT = Path(sys.executable).resolve().parent if getattr(sys, "frozen", False) else Path(__file__).resolve().parents[1]
PROJECT_TMP = PROJECT_ROOT / "work" / "tmp"
GROUP_CONFIG_PATH = PROJECT_ROOT / "docs" / "平谷区街乡镇分组清单.json"
WEEKLY_LOGIC_PATH = PROJECT_ROOT / "docs" / "周报生成逻辑汇总.md"

CITY_CATEGORIES = [
    "居民投放情况",
    "值守指导",
    "投放点情况",
    "大件垃圾投放点",
    "规范管理",
    "小区公示牌",
]

GROUP_ORDER = ["城区组", "平原组", "山区组"]
DEFAULT_GROUP_CONFIG = {
    "城区组": ["滨河街道", "平谷镇", "兴谷街道"],
    "平原组": [
        "山东庄镇",
        "马坊镇",
        "马昌营镇",
        "金海湖镇",
        "南独乐河镇",
        "峪口镇",
        "东高村镇",
        "大兴庄镇",
        "王辛庄镇",
        "夏各庄镇",
        "黄松峪乡",
    ],
    "山区组": ["镇罗营镇", "刘家店镇", "熊儿寨乡", "大华山镇"],
}
CN_INDEX = ["一", "二", "三", "四", "五", "六"]
NO_ISSUE_RE = re.compile(r"(没问题|无问题|无异常|查\s*\d+\s*错\s*0|查\d+错0)")
DISTRICT_TOP_EXCLUDED_ISSUES = {"容器完好率"}
MERGED_CONTAINER_ISSUES = {"容器满冒", "容器整洁率"}
MERGED_CONTAINER_ISSUE_NAME = "容器满冒脏污"
ISSUE_PHRASE_MAP = {
    "居民自主投放": "居民自主投放不规范",
    "桶站及周边环境": "桶站及周边环境不整洁",
    "容器满冒": "分类容器满冒",
    "容器整洁率": "分类容器脏污",
    "容器满冒脏污": "容器满冒脏污",
    "容器完好率": "分类容器破损",
    "垃圾分类驿站": "生活垃圾分类驿站问题",
    "大件垃圾投放点设置": "大件垃圾投放点设置不规范",
    "装修垃圾投放点设置": "装修垃圾投放点设置不规范",
    "桶站便利措施": "桶站便利措施不完善",
    "公示牌": "公示牌信息不规范",
    "分类标识": "分类标识不规范",
    "宣传氛围": "垃圾分类宣传氛围不足",
    "责任公示": "管理责任公示不规范",
    "垃圾桶成组配备": "分类容器未成组配备",
    "投放点设置": "投放点设置不规范",
    "大件垃圾管理": "大件垃圾管理不规范",
    "其他": "其他垃圾分类管理问题",
}


class WeeklyReportError(Exception):
    """User-facing validation or generation error."""


@dataclass
class GenerationResult:
    report_path: Path
    context: dict[str, Any]


def ensure_project_temp() -> Path:
    PROJECT_TMP.mkdir(parents=True, exist_ok=True)
    for key in ("TMP", "TEMP", "TMPDIR"):
        os.environ[key] = str(PROJECT_TMP)
    return PROJECT_TMP


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def to_number(value: Any) -> int:
    if value is None or value == "":
        return 0
    try:
        return int(float(value))
    except (TypeError, ValueError):
        nums = re.findall(r"\d+", str(value))
        return int(nums[0]) if nums else 0


def pct(part: int | float, total: int | float, digits: int = 1) -> str:
    if not total:
        return f"{0:.{digits}f}"
    return f"{part / total * 100:.{digits}f}"


def normalize_date(value: Any) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, (int, float)):
        digits = str(int(value))
        if len(digits) == 8:
            return f"{digits[:4]}-{digits[4:6]}-{digits[6:8]}"
        try:
            base = datetime(1899, 12, 30)
            return (base + timedelta(days=int(value))).strftime("%Y-%m-%d")
        except Exception:
            return digits
    text = str(value).strip()
    match = re.search(r"(20\d{2})[./年-]?(\d{1,2})[./月-]?(\d{1,2})", text)
    if match:
        y, m, d = match.groups()
        return f"{int(y):04d}-{int(m):02d}-{int(d):02d}"
    return text


def parse_datetime_value(value: Any) -> datetime | None:
    if value in (None, ""):
        return None
    if isinstance(value, datetime):
        return value
    if isinstance(value, (int, float)):
        try:
            return datetime(1899, 12, 30) + timedelta(days=float(value))
        except Exception:
            return None
    text = str(value).strip()
    if not text:
        return None
    patterns = [
        r"(20\d{2})[-/.年](\d{1,2})[-/.月](\d{1,2})日?\s+(\d{1,2}):(\d{1,2})(?::(\d{1,2}))?",
        r"(20\d{2})[-/.年](\d{1,2})[-/.月](\d{1,2})日?",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if not match:
            continue
        parts = [int(p) if p else 0 for p in match.groups()]
        while len(parts) < 6:
            parts.append(0)
        try:
            return datetime(parts[0], parts[1], parts[2], parts[3], parts[4], parts[5])
        except ValueError:
            return None
    return None


def case_date_from_row(row: dict[str, Any]) -> datetime | None:
    check_candidates = [
        row.get("检查时间"),
        row.get("检查日期"),
        row.get("具体问题_11"),
    ]
    for value in check_candidates:
        parsed = parse_datetime_value(value)
        if parsed is not None:
            return datetime(parsed.year, parsed.month, parsed.day)
    report_time = parse_datetime_value(row.get("上报时间"))
    if report_time is None:
        return None
    case_date = report_time.date()
    if report_time.hour >= 12:
        case_date = case_date + timedelta(days=1)
    return datetime(case_date.year, case_date.month, case_date.day)


def parse_period_dates(period_short: str) -> tuple[datetime, datetime] | None:
    match = re.fullmatch(r"(\d{2})\.(\d{2})-(\d{2})\.(\d{2})", period_short or "")
    if not match:
        return None
    current_year = datetime.now().year
    sm, sd, em, ed = [int(part) for part in match.groups()]
    start = datetime(current_year, sm, sd)
    end = datetime(current_year, em, ed)
    if end < start:
        end = datetime(current_year + 1, em, ed)
    return start, end


def filter_rows_by_case_period(rows: list[dict[str, Any]], period_short: str) -> list[dict[str, Any]]:
    bounds = parse_period_dates(period_short)
    if bounds is None:
        return rows
    start, end = bounds
    filtered = []
    for row in rows:
        case_date = row.get("案件日期")
        if isinstance(case_date, datetime) and start <= case_date <= end:
            filtered.append(row)
        elif not isinstance(case_date, datetime):
            filtered.append(row)
    return filtered


def date_value_as_datetime(value: Any) -> datetime | None:
    normalized = normalize_date(value)
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", normalized):
        return None
    try:
        return datetime.strptime(normalized, "%Y-%m-%d")
    except ValueError:
        return None


def filter_city_rows_by_period(rows: list[dict[str, Any]], period_short: str) -> list[dict[str, Any]]:
    bounds = parse_period_dates(period_short)
    if bounds is None:
        return rows
    start, end = bounds
    return [
        row
        for row in rows
        if (row_date := date_value_as_datetime(row.get("日期"))) is not None and start <= row_date <= end
    ]


def short_period_from_dates(dates: Iterable[Any]) -> str:
    normalized = [normalize_date(v) for v in dates]
    normalized = [v for v in normalized if re.match(r"\d{4}-\d{2}-\d{2}", v)]
    if not normalized:
        return ""
    start, end = min(normalized), max(normalized)
    return f"{start[5:7]}.{start[8:10]}-{end[5:7]}.{end[8:10]}"


def parse_period_short(text: str, fallback_dates: Iterable[Any] = ()) -> str:
    text = (text or "").strip()
    if text:
        matches = re.findall(r"(\d{1,2})[.月/-](\d{1,2})", text)
        if len(matches) >= 2:
            a, b = matches[0], matches[-1]
            return f"{int(a[0]):02d}.{int(a[1]):02d}-{int(b[0]):02d}.{int(b[1]):02d}"
        matches = re.findall(r"(\d{1,2})\s*月\s*(\d{1,2})\s*日", text)
        if len(matches) >= 2:
            a, b = matches[0], matches[-1]
            return f"{int(a[0]):02d}.{int(a[1]):02d}-{int(b[0]):02d}.{int(b[1]):02d}"
    return short_period_from_dates(fallback_dates)


def period_for_filename(period_short: str) -> str:
    return period_short.replace(".", ".")


def report_filename(report_year: int, period_short: str) -> str:
    period_name = period_for_filename(period_short)
    return f"{report_year}年平谷区垃圾分类桶站检查情况通报（{period_name}）.docx"


def period_display(period_short: str) -> str:
    match = re.fullmatch(r"(\d{2})\.(\d{2})-(\d{2})\.(\d{2})", period_short or "")
    if not match:
        return period_short
    sm, sd, em, ed = match.groups()
    return f"{sm}月{sd}日-{em}月{ed}日"


def issue_comparison_text(current: int, previous: int | None) -> str:
    if previous is None:
        return ""
    diff = current - previous
    if previous == 0:
        if diff == 0:
            return "较上周持平。"
        return f"较上周{previous}处增加{diff}处。"
    rate = abs(diff) / previous * 100
    if diff < 0:
        return f"较上周{previous}处减少{abs(diff)}处，降幅{rate:.1f}%。"
    if diff > 0:
        return f"较上周{previous}处增加{diff}处，增幅{rate:.1f}%。"
    return f"较上周{previous}处持平。"


def sanitize_problem_text(text: Any) -> str:
    value = clean_text(text)
    value = re.sub(r"投放点\s*\d+\s*[：:、,，.．]\s*", "", value)
    value = re.sub(r"(?:(?<=^)|(?<=[；;，,。]))\s*\d+\s*[：:、,，.．]\s*", "", value)
    value = re.sub(r"\s+", "", value)
    value = value.strip("；;，,。 ")
    return value


def problem_label(text: Any) -> str:
    value = sanitize_problem_text(text)
    value = re.sub(r"\d+\s*(处|个|人)$", "", value)
    value = value.strip("；;，,。 ")
    replacements = [
        ("其他垃圾分类投放不准确", "其他垃圾投放不准确"),
        ("居民自主投放分类投放不准确", "居民自主投放不准确"),
        ("厢房式、封闭式投放点外部周边摆放垃圾桶", "周边摆放垃圾桶"),
        ("散落垃圾等环境脏乱", "环境脏乱"),
        ("厨余、其他、可回收物垃圾桶未成组配备", "垃圾桶未成组配备"),
        ("居民不知晓垃圾分类，参与垃圾分类，便民回收渠道，开展垃圾分类宣传活动", "居民不知晓垃圾分类"),
    ]
    for old, new in replacements:
        value = value.replace(old, new)
    value = re.sub(r"居民不知晓垃圾分类.*", "居民不知晓垃圾分类", value)
    value = re.sub(r"无大件垃圾托底上门回收信息.*", "无大件垃圾托底上门回收信息", value)
    value = re.sub(r"未配置灭火器材或配置但不完好.*", "未配置灭火器材或配置但不完好", value)
    value = re.sub(r"标志不合格.*", "标志不合格", value)
    value = re.sub(r"大件装修投放点信息不准确.*", "大件装修投放点信息不准确", value)
    value = re.sub(r"有明显散落垃圾.*", "有明显散落垃圾", value)
    value = re.sub(r"投放点有满冒.*", "投放点有满冒", value)
    return value


def problem_labels(text: Any) -> list[str]:
    value = sanitize_problem_text(text)
    parts = re.split(r"[；;]", value)
    expanded: list[str] = []
    for part in parts:
        if "," in part:
            expanded.extend(part.split(","))
        else:
            expanded.append(part)
    labels = []
    for part in expanded:
        label = problem_label(part)
        if label and label not in labels:
            labels.append(label)
    fallback = problem_label(value)
    return labels or ([fallback] if fallback else [])


def join_names(names: Iterable[str], limit: int | None = None) -> str:
    values = [name for name in names if name]
    if limit is not None:
        values = values[:limit]
    return "、".join(values)


def issue_phrase(issue_name: str) -> str:
    return ISSUE_PHRASE_MAP.get(issue_name, f"{issue_name}问题")


def read_xls_rows(path: Path) -> tuple[list[str], list[dict[str, Any]]]:
    book = xlrd.open_workbook(str(path), on_demand=True)
    try:
        sheet = book.sheet_by_index(0)
        if sheet.nrows < 3:
            raise WeeklyReportError("区级基础台账行数不足，至少需要两行表头和数据行。")
        headers = [clean_text(v) for v in sheet.row_values(1)]
        rows: list[dict[str, Any]] = []
        for r in range(2, sheet.nrows):
            values = sheet.row_values(r)
            if not any(clean_text(v) for v in values):
                continue
            item = {}
            for idx, header in enumerate(headers):
                key = header or f"空列{idx + 1}"
                if key in item:
                    key = f"{key}_{idx + 1}"
                item[key] = values[idx] if idx < len(values) else ""
            rows.append(item)
        return headers, rows
    finally:
        book.release_resources()


def required_columns(headers: Iterable[str], names: Iterable[str], label: str) -> None:
    header_set = set(headers)
    missing = [name for name in names if name not in header_set]
    if missing:
        raise WeeklyReportError(f"{label}缺少必要列：{', '.join(missing)}")


def is_issue_row(row: dict[str, Any]) -> bool:
    issue = clean_text(row.get("具体问题"))
    second = clean_text(row.get("2级指标"))
    third = clean_text(row.get("3级指标"))
    merged = f"{issue} {second} {third}"
    return bool(issue) and not NO_ISSUE_RE.search(merged)


def read_district_base(path: Path) -> dict[str, Any]:
    headers, rows = read_xls_rows(path)
    required_columns(headers, ["2级点位", "3级点位", "5级点位", "具体问题", "2级指标", "3级指标"], "区级基础台账")
    for row in rows:
        case_date = case_date_from_row(row)
        row["案件日期"] = case_date
        if case_date is not None and not clean_text(row.get("检查时间")):
            row["检查时间"] = case_date.strftime("%Y-%m-%d")
    issue_rows = [row for row in rows if is_issue_row(row)]
    return {
        "headers": headers,
        "rows": rows,
        "issue_rows": issue_rows,
        "period_short": parse_period_short(path.stem, [r.get("案件日期") or r.get("上报时间") for r in rows]),
    }


def default_city_sheet(path: Path) -> str:
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    try:
        candidates = []
        for ws in wb.worksheets:
            count = 0
            for row in ws.iter_rows(values_only=True):
                if any(clean_text(v) for v in row):
                    count += 1
            if count > 1:
                candidates.append(ws.title)
        if not candidates:
            raise WeeklyReportError("市级台账没有识别到含数据的工作表。")
        return candidates[-1]
    finally:
        wb.close()


def city_sheet_names(path: Path) -> list[str]:
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    try:
        return wb.sheetnames
    finally:
        wb.close()


def read_city_ledger(path: Path, sheet_name: str | None = None) -> dict[str, Any]:
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    try:
        if sheet_name is None:
            sheet_name = default_city_sheet(path)
        if sheet_name not in wb.sheetnames:
            raise WeeklyReportError(f"市级台账中不存在工作表：{sheet_name}")
        ws = wb[sheet_name]
        rows_iter = ws.iter_rows(values_only=True)
        header_values = next(rows_iter, None)
        if not header_values:
            raise WeeklyReportError("市级台账工作表为空。")
        headers = [clean_text(v) for v in header_values]
        col_map = {
            "日期": _index(headers, "日期"),
            "乡镇街道": _index(headers, "乡镇街道"),
            "小区/村": _index(headers, "小区/村"),
            "大类问题": _index(headers, "大类问题"),
            "问题": _index(headers, "问题"),
            "问题数量": 6,
        }
        for key, idx in col_map.items():
            if idx is None or idx >= len(headers):
                raise WeeklyReportError(f"市级台账缺少必要列：{key}")
        rows = []
        for values in rows_iter:
            if not any(clean_text(v) for v in values):
                continue
            rows.append(
                {
                    "日期": values[col_map["日期"]],
                    "乡镇街道": clean_text(values[col_map["乡镇街道"]]),
                    "小区/村": clean_text(values[col_map["小区/村"]]),
                    "大类问题": clean_text(values[col_map["大类问题"]]),
                    "问题": sanitize_problem_text(values[col_map["问题"]]),
                    "问题数量": to_number(values[col_map["问题数量"]]),
                }
            )
        return {
            "sheet_name": sheet_name,
            "headers": headers,
            "rows": rows,
            "period_short": parse_period_short(path.stem, [r["日期"] for r in rows]),
        }
    finally:
        wb.close()


def apply_city_period_filter(city: dict[str, Any], period_short: str) -> None:
    source_rows = city["rows"]
    data_period_short = short_period_from_dates(row.get("日期") for row in source_rows)
    filtered_rows = filter_city_rows_by_period(source_rows, period_short)
    if not filtered_rows:
        data_period_text = data_period_short or "无法识别"
        raise WeeklyReportError(
            f"市级台账工作表“{city['sheet_name']}”在周期 {period_short} 内没有匹配记录"
            f"（该表实际日期范围：{data_period_text}），请重新选择市级工作表或修正市级周期。"
        )
    city["source_row_count"] = len(source_rows)
    city["filtered_out_count"] = len(source_rows) - len(filtered_rows)
    city["data_period_short"] = data_period_short
    city["rows"] = filtered_rows


def _index(headers: list[str], name: str) -> int | None:
    try:
        return headers.index(name)
    except ValueError:
        return None


def ensure_group_config_file() -> None:
    if GROUP_CONFIG_PATH.exists():
        return
    GROUP_CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    GROUP_CONFIG_PATH.write_text(json.dumps(DEFAULT_GROUP_CONFIG, ensure_ascii=False, indent=2), encoding="utf-8")


def groups_config_to_mapping(config: dict[str, Any]) -> dict[str, str]:
    result: dict[str, str] = {}
    for group, towns in config.items():
        if not isinstance(towns, list):
            raise WeeklyReportError(f"分组 JSON 中“{group}”应为街乡镇列表。")
        for town in towns:
            town_name = clean_text(town)
            if town_name:
                result[town_name] = clean_text(group)
    return result


def groups_mapping_to_config(groups: dict[str, str]) -> dict[str, list[str]]:
    config: dict[str, list[str]] = {group: [] for group in GROUP_ORDER}
    for town, group in groups.items():
        config.setdefault(group, [])
        if town not in config[group]:
            config[group].append(town)
    ordered: dict[str, list[str]] = {}
    for group in GROUP_ORDER:
        if group in config:
            ordered[group] = config[group]
    for group in sorted(g for g in config if g not in ordered):
        ordered[group] = config[group]
    return ordered


def load_groups_json(path: Path = GROUP_CONFIG_PATH) -> dict[str, str]:
    ensure_group_config_file()
    try:
        config = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise WeeklyReportError(f"分组 JSON 格式错误：{exc}") from exc
    if not isinstance(config, dict):
        raise WeeklyReportError("分组 JSON 顶层必须是对象。")
    return groups_config_to_mapping(config)


def save_groups_json(groups: dict[str, str], path: Path = GROUP_CONFIG_PATH) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    config = groups_mapping_to_config(groups)
    path.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")


def add_group_point(town: str, group: str, path: Path = GROUP_CONFIG_PATH) -> dict[str, str]:
    town_name = clean_text(town)
    group_name = clean_text(group)
    if not town_name:
        raise WeeklyReportError("请填写城乡镇。")
    if not group_name:
        raise WeeklyReportError("请填写分组。")
    groups = load_groups_json(path)
    groups[town_name] = group_name
    save_groups_json(groups, path)
    return groups


def read_groups_excel(path: Path) -> dict[str, str]:
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb.active
        header = [clean_text(v) for v in next(ws.iter_rows(values_only=True))]
        required_columns(header, ["分组", "街乡镇"], "街乡镇分组清单")
        group_idx, town_idx = header.index("分组"), header.index("街乡镇")
        current_group = ""
        result = {}
        for row in ws.iter_rows(min_row=2, values_only=True):
            group = clean_text(row[group_idx])
            town = clean_text(row[town_idx])
            if group:
                current_group = group
            if town:
                result[town] = current_group
        return result
    finally:
        wb.close()


def read_groups(path: Path) -> dict[str, str]:
    if path.suffix.lower() == ".json":
        return load_groups_json(path)
    return read_groups_excel(path)


def resolve_groups(group_path: Path | None = None) -> tuple[dict[str, str], str]:
    if group_path is not None and str(group_path).strip():
        groups = read_groups_excel(group_path)
        current = load_groups_json()
        if groups != current:
            save_groups_json(groups)
        return groups, f"已使用上传分组并同步 JSON：{group_path.name}"
    groups = load_groups_json()
    return groups, f"已使用固定 JSON 分组：{GROUP_CONFIG_PATH.name}"


def city_category(raw: str) -> str:
    raw = raw or ""
    if "公示牌" in raw:
        return "小区公示牌"
    if "大件垃圾投放点" in raw:
        return "大件垃圾投放点"
    if "值守" in raw:
        return "值守指导"
    if "投放点情况" in raw:
        return "投放点情况"
    if "居民询问情况" in raw or "环境脏乱" in raw:
        return "规范管理"
    if "居民投放" in raw or "居民自主投放" in raw:
        return "居民投放情况"
    return "规范管理"


def top_items(counter: Counter, limit: int = 3) -> list[tuple[str, int]]:
    return [(k, v) for k, v in counter.most_common(limit) if k and v]


def district_issue_bucket(value: Any) -> str:
    name = clean_text(value)
    if name in MERGED_CONTAINER_ISSUES:
        return MERGED_CONTAINER_ISSUE_NAME
    return name


def join_counts(items: Iterable[tuple[str, int]], suffix: str = "处") -> str:
    return "、".join(f"{name}{count}{suffix}" for name, count in items)


def checked_point_counts(rows: list[dict[str, Any]]) -> tuple[int, int, dict[str, Counter]]:
    community: set[tuple[str, str]] = set()
    village: set[tuple[str, str]] = set()
    by_town = {"小区": Counter(), "村居": Counter()}
    for row in rows:
        town = clean_text(row.get("3级点位"))
        level2 = clean_text(row.get("2级点位"))
        site = clean_text(row.get("5级点位"))
        if not town or not site:
            continue
        if level2 == "小区":
            community.add((town, site))
            by_town["小区"][(town, site)] = 1
        elif level2 == "村居":
            village.add((town, site))
            by_town["村居"][(town, site)] = 1
    town_counts = {"小区": Counter(), "村居": Counter()}
    for label in ("小区", "村居"):
        for town, _site in by_town[label]:
            town_counts[label][town] += 1
    return len(community), len(village), town_counts


def all_checked_counts(rows: list[dict[str, Any]]) -> Counter:
    points: set[tuple[str, str, str]] = set()
    for row in rows:
        town = clean_text(row.get("3级点位"))
        level2 = clean_text(row.get("2级点位"))
        site = clean_text(row.get("5级点位"))
        if town and site:
            points.add((town, site, level2))
    counts = Counter()
    for town, _site, _level2 in points:
        counts[town] += 1
    return counts


def build_context(
    district: dict[str, Any],
    city: dict[str, Any],
    groups: dict[str, str],
    template_path: Path,
    report_year: str,
    district_period_short: str,
    city_period_short: str,
    previous_district_issues: int | None = None,
) -> dict[str, Any]:
    district_rows = district["rows"]
    district_issue_rows = district["issue_rows"]
    city_rows = city["rows"]

    district_towns = {clean_text(r.get("3级点位")) for r in district_rows if clean_text(r.get("3级点位"))}
    missing_groups = sorted(t for t in district_towns if t not in groups)
    if missing_groups:
        raise WeeklyReportError("分组清单缺少以下乡镇街道：" + "、".join(missing_groups))

    city_towns = {r["乡镇街道"] for r in city_rows if r["乡镇街道"]}
    city_points = {(r["乡镇街道"], r["小区/村"]) for r in city_rows if r["乡镇街道"] and r["小区/村"]}
    city_total = sum(r["问题数量"] for r in city_rows)
    city_point_by_town = Counter(town for town, _site in city_points)
    city_period_text = period_display(city_period_short)
    district_period_text = period_display(district_period_short)
    city_overview = (
        f"{report_year}年{city_period_text}，市级检查了{len(city_towns)}个乡镇"
        f"{len(city_points)}个小区（{join_counts(city_point_by_town.most_common(), '个')}）的桶站垃圾分类情况，"
        f"共发现{city_total}处桶站垃圾分类问题。"
    )

    community_count, village_count, district_town_counts = checked_point_counts(district_rows)
    community_desc = join_counts(district_town_counts["小区"].most_common(), "个")
    village_desc = join_counts(district_town_counts["村居"].most_common(), "个")
    comparison = issue_comparison_text(len(district_issue_rows), previous_district_issues)
    comparison = f"{comparison}" if comparison else ""
    district_overview = (
        f"{report_year}年{district_period_text}，区级检查了{len(district_towns)}个乡镇（街道）"
        f"{community_count}个小区（{community_desc}）和{village_count}个村居（{village_desc}）的桶站垃圾分类情况，"
        f"共发现{len(district_issue_rows)}处桶站垃圾分类问题。{comparison}"
    )

    city_category_rows, city_category_paragraphs = city_six_category_analysis(city_rows, city_total)
    concentrated_summary, concentrated_paragraphs = city_concentrated_analysis(city_rows, city_total)
    rank_rows = district_rank_analysis(district_rows, district_issue_rows, groups)
    district_summary = district_issue_analysis(district_issue_rows, groups)
    top_town_paragraphs = district_top_town_analysis(rank_rows, district_issue_rows)
    next_work = next_work_items(city_rows, district_issue_rows)
    warnings = work_action_warnings()

    context = {
        "report_year": report_year,
        "city_period_short": city_period_short,
        "district_period_short": district_period_short,
        "city_period_display": city_period_short,
        "district_period_display": district_period_short,
        "city_period_body": city_period_text,
        "district_period_body": district_period_text,
        "city_overview": city_overview,
        "district_overview": district_overview,
        "city_six_categories": city_category_rows,
        "city_six_category_paragraphs": city_category_paragraphs,
        "city_concentrated_area_summary": concentrated_summary,
        "city_concentrated_area_paragraphs": concentrated_paragraphs,
        "district_issue_summary": district_summary,
        "district_top_town_paragraphs": top_town_paragraphs,
        "district_rank_rows": rank_rows,
        "next_work_items": next_work,
        "warnings": warnings,
    }
    validate_template_variables(template_path, context)
    return context


def city_six_category_analysis(rows: list[dict[str, Any]], total: int) -> tuple[list[dict[str, Any]], list[str]]:
    counts = Counter()
    problems: dict[str, Counter] = defaultdict(Counter)
    towns: dict[str, Counter] = defaultdict(Counter)
    for row in rows:
        category = city_category(row["大类问题"])
        count = row["问题数量"]
        counts[category] += count
        for label in problem_labels(row["问题"]):
            problems[category][label] += count
        towns[category][row["乡镇街道"]] += count

    table_rows = []
    paragraphs = []
    for idx, name in enumerate(CITY_CATEGORIES, start=1):
        count = counts[name]
        main = top_items(problems[name], 4)
        town_top = top_items(towns[name], 3)
        if count == 0:
            desc = "经上周期问题整改后持续保持良好状态，本周期未发现相关问题"
            paragraph = f"{idx}.{short_city_title(name)}。经上周期问题整改后持续保持良好状态，本周期未发现相关问题。"
        else:
            main_names = [item[0] for item in main]
            main_text = f"主要问题为{join_names(main_names, 3)}" if main_names else ""
            town_text = f"问题突出的乡镇是{join_counts(town_top)}" if town_top else ""
            desc = city_table_description(name, main, town_top)
            paragraph = (
                f"{idx}.{short_city_title(name)}。本周期发现{count}处，占问题总数的{pct(count, total)}%，"
                f"{main_text + '，' if main_text else ''}{town_text}。"
            )
        table_rows.append({"name": city_table_name(name), "count": count, "description": desc})
        paragraphs.append(paragraph)
    return table_rows, paragraphs


def city_table_name(name: str) -> str:
    return "投放规范" if name == "规范管理" else name


def city_table_description(category: str, main: list[tuple[str, int]], towns: list[tuple[str, int]]) -> str:
    main_names = [name for name, _count in main if name]
    town_names = [name for name, _count in towns if name]
    top_problem = main_names[0] if main_names else "相关问题"
    top_two_towns = "、".join(town_names[:2])
    if category == "居民投放情况":
        return f"“{top_problem}”为最突出问题，在多小区反复出现，{top_two_towns}问题集中"
    if category == "值守指导":
        return "上周期问题整改后持续保持良好状态，本周期未发现相关问题"
    if category == "投放点情况":
        return f"{join_names(main_names, 3)}为主要问题，{top_two_towns}需强化设施管理"
    if category == "大件垃圾投放点":
        return f"主要集中在{top_two_towns}，{join_names(main_names, 2)}问题需整改"
    if category == "规范管理":
        return f"{join_names(main_names, 2)}问题出现反弹，{top_two_towns}仍需加强管理"
    if category == "小区公示牌":
        return f"主要集中在{top_two_towns}，{join_names(main_names, 4)}为主要问题"
    return f"{join_names(main_names, 3)}为主要问题，{top_two_towns}问题集中"


def short_city_title(name: str) -> str:
    return {
        "居民投放情况": "居民投放问题",
        "值守指导": "人员值守问题",
        "投放点情况": "投放点问题",
        "大件垃圾投放点": "大件垃圾问题",
        "规范管理": "规范管理问题",
        "小区公示牌": "小区公示牌类问题",
    }[name]


def city_concentrated_analysis(rows: list[dict[str, Any]], total: int) -> tuple[str, list[str]]:
    town_counts = Counter()
    town_cat_counts: dict[str, Counter] = defaultdict(Counter)
    for row in rows:
        town = row["乡镇街道"]
        count = row["问题数量"]
        town_counts[town] += count
        town_cat_counts[town][city_category(row["大类问题"])] += count
    top = top_items(town_counts, 3)
    if not top:
        return "根据本周期市级检查数据统计，未发现市级问题集中区域。", []
    top_sum = sum(v for _k, v in top)
    summary = f"根据本周期市级检查数据统计，问题数量较多的乡镇（街道）为{join_counts(top)}"
    if total and top_sum / total >= 0.6:
        summary += f"，以上{len(top)}个乡镇占市级总问题数{pct(top_sum, total)}%"
    summary += "，核心问题如下。"
    paragraphs = []
    for idx, (town, count) in enumerate(top, start=1):
        cat_text = join_counts(top_items(town_cat_counts[town], 6))
        suffix = "，问题最多" if idx == 1 else ""
        paragraphs.append(f"{idx}.{town}共发现问题{count}处{suffix}。{cat_text}。")
    return summary, paragraphs


def district_rank_analysis(
    district_rows: list[dict[str, Any]],
    issue_rows: list[dict[str, Any]],
    groups: dict[str, str],
) -> list[dict[str, Any]]:
    checked = all_checked_counts(district_rows)
    issues = Counter(clean_text(row.get("3级点位")) for row in issue_rows)
    rows = []
    for town, checked_count in checked.items():
        group = groups.get(town, "")
        issue_count = issues[town]
        rate = issue_count / checked_count if checked_count else 0
        rows.append(
            {
                "group": group,
                "town": town,
                "issue_count": issue_count,
                "checked_count": checked_count,
                "issue_rate_value": rate,
                "issue_rate": f"{rate:.2f}",
                "rank": 0,
            }
        )
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        grouped[row["group"]].append(row)
    for group_rows in grouped.values():
        for rank, row in enumerate(sorted(group_rows, key=lambda r: (r["issue_rate_value"], r["town"])), start=1):
            row["rank"] = rank
    ordered = []
    for group in GROUP_ORDER:
        ordered.extend(sorted(grouped.get(group, []), key=lambda r: (-r["issue_rate_value"], r["town"])))
    other_groups = sorted(g for g in grouped if g not in GROUP_ORDER)
    for group in other_groups:
        ordered.extend(sorted(grouped[group], key=lambda r: (-r["issue_rate_value"], r["town"])))
    last_group = None
    for row in ordered:
        row["group_first"] = row["group"] != last_group
        last_group = row["group"]
        row.pop("issue_rate_value", None)
    return ordered


def district_issue_analysis(issue_rows: list[dict[str, Any]], groups: dict[str, str]) -> str:
    total = len(issue_rows)
    issue_counts = Counter(district_issue_bucket(row.get("2级指标")) for row in issue_rows)
    top3 = top_items(issue_counts, 3)
    top_text = "、".join(f"{name}{count}处（占比{pct(count, total, 2)}%）" for name, count in top3)
    group_parts = []
    by_group: dict[str, Counter] = defaultdict(Counter)
    for row in issue_rows:
        town = clean_text(row.get("3级点位"))
        group = groups.get(town, "")
        by_group[group][district_issue_bucket(row.get("2级指标"))] += 1
    for group in GROUP_ORDER:
        issues = top_items(by_group[group], 2)
        if issues:
            group_parts.append(f"{group}主要问题为{'、'.join(name + '问题' for name, _count in issues)}")
    return f"本周期累计发现问题{total}处，突出问题为{top_text}。其中，{'，'.join(group_parts)}。"


def district_top_town_analysis(rank_rows: list[dict[str, Any]], issue_rows: list[dict[str, Any]]) -> list[str]:
    by_group = defaultdict(list)
    for row in rank_rows:
        by_group[row["group"]].append(row)
    all_top = max(rank_rows, key=lambda r: (float(r["issue_rate"]), r["issue_count"]), default=None)
    issues_by_town: dict[str, Counter] = defaultdict(Counter)
    for row in issue_rows:
        issue_name = district_issue_bucket(row.get("2级指标"))
        if issue_name in DISTRICT_TOP_EXCLUDED_ISSUES:
            continue
        issues_by_town[clean_text(row.get("3级点位"))][issue_name] += 1
    paragraphs = []
    seq = 1
    for group in GROUP_ORDER:
        candidates = by_group.get(group, [])
        if not candidates:
            continue
        top = max(candidates, key=lambda r: (float(r["issue_rate"]), r["issue_count"]))
        main = "、".join(f"{issue_phrase(name)}（{count}处）" for name, count in top_items(issues_by_town[top["town"]], 4))
        if all_top and top["town"] == all_top["town"]:
            rank_text = "位列本周检查乡镇问题率首位"
        else:
            rank_text = f"位列{group}问题率首位"
        paragraphs.append(
            f"{seq}.{top['town']}共检查{top['checked_count']}个小区（村），发现问题{top['issue_count']}处，"
            f"问题率{top['issue_rate']}%，{rank_text}，主要问题为{main}。"
        )
        seq += 1
    return paragraphs


WORK_TITLES = {
    "居民投放情况": "强化源头分类引导，提升居民自主投放准确率",
    "值守指导": "压实桶前值守责任，提升现场分类指导质效",
    "投放点情况": "加强桶站设施管理，改善投放环境面貌",
    "大件垃圾投放点": "规范大件垃圾及驿站管理，补齐设施短板",
    "规范管理": "健全长效管理机制，压实各方工作责任",
    "小区公示牌": "规范公示信息管理，提升信息准确性",
}

def split_markdown_row(line: str) -> list[str]:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return cells


def split_action_keywords(text: str) -> tuple[str, ...]:
    return tuple(
        item.strip()
        for item in re.split(r"[、，,；;]+", text)
        if item.strip() and item.strip() != "默认"
    )


def fallback_work_action(category: str) -> str:
    title = WORK_TITLES.get(category, "相关工作")
    return f"围绕“{title}”持续开展问题排查和整改，压实责任、跟踪销号，推动相关问题及时整改到位。"


@lru_cache(maxsize=1)
def load_work_action_config() -> tuple[dict[str, list[tuple[tuple[str, ...], str]]], dict[str, str], tuple[str, ...]]:
    warnings: list[str] = []
    if not WEEKLY_LOGIC_PATH.exists():
        warnings.append(f"缺少周报生成逻辑文档：{WEEKLY_LOGIC_PATH}；下一步重点工作已使用通用兜底句。")
        return {}, {name: fallback_work_action(name) for name in WORK_TITLES}, tuple(warnings)
    try:
        lines = WEEKLY_LOGIC_PATH.read_text(encoding="utf-8").splitlines()
    except OSError as exc:
        warnings.append(f"读取周报生成逻辑文档失败：{exc}；下一步重点工作已使用通用兜底句。")
        return {}, {name: fallback_work_action(name) for name in WORK_TITLES}, tuple(warnings)
    in_section = False
    table_started = False
    rules: dict[str, list[tuple[tuple[str, ...], str]]] = defaultdict(list)
    defaults: dict[str, str] = {}
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("主要措施句匹配口径"):
            in_section = True
            continue
        if not in_section:
            continue
        if stripped.startswith("## "):
            break
        if not stripped.startswith("|"):
            continue
        cells = split_markdown_row(stripped)
        if len(cells) < 4:
            continue
        if cells[0] == "工作方向":
            table_started = True
            continue
        if not table_started or set(cells[0]) <= {"-", ":"}:
            continue
        category, match_text, sentence = cells[0], cells[1], cells[2]
        if category not in WORK_TITLES or not sentence:
            continue
        if match_text == "默认":
            defaults[category] = sentence
            continue
        keywords = split_action_keywords(match_text)
        if keywords:
            rules[category].append((keywords, sentence))
    missing_defaults = [name for name in WORK_TITLES if name not in defaults]
    if missing_defaults:
        warnings.append("周报生成逻辑文档缺少默认措施句：" + "、".join(missing_defaults) + "；对应方向已使用通用兜底句。")
        for name in missing_defaults:
            defaults[name] = fallback_work_action(name)
    if not table_started:
        warnings.append("周报生成逻辑文档未找到“主要措施句匹配口径”表；下一步重点工作已使用通用兜底句。")
    return dict(rules), defaults, tuple(warnings)


def work_action_warnings() -> list[str]:
    _rules, _defaults, warnings = load_work_action_config()
    return list(warnings)


def next_work_items(city_rows: list[dict[str, Any]], issue_rows: list[dict[str, Any]]) -> list[dict[str, str]]:
    counts = Counter()
    city_problems: dict[str, Counter] = defaultdict(Counter)
    for row in city_rows:
        category = city_category(row["大类问题"])
        count = row["问题数量"]
        counts[category] += count
        if count <= 0:
            continue
        for label in problem_labels(row["问题"]):
            city_problems[category][label] += count
    selected = [name for name, count in counts.most_common() if count > 0][:4]
    district_problems: dict[str, Counter] = defaultdict(Counter)
    for row in issue_rows:
        issue_name = district_issue_bucket(row.get("2级指标"))
        category = work_category_for_district_issue(issue_name)
        if category:
            district_problems[category][issue_name] += 1
    district_top = [name for name, _count in top_items(Counter(district_issue_bucket(r.get("2级指标")) for r in issue_rows), 3)]
    if any(any(key in issue for key in [MERGED_CONTAINER_ISSUE_NAME, "桶站及周边环境", "垃圾分类驿站"]) for issue in district_top):
        if "投放点情况" not in selected:
            selected = (selected[:3] + ["投放点情况"]) if len(selected) >= 4 else selected + ["投放点情况"]
    items = []
    for idx, name in enumerate(selected[:4]):
        items.append(
            {
                "cn_index": CN_INDEX[idx],
                "title": WORK_TITLES[name],
                "body": work_action_body(name, city_problems[name], district_problems[name]),
            }
        )
    return items


def work_category_for_district_issue(issue_name: str) -> str:
    if issue_name in {"居民自主投放"}:
        return "居民投放情况"
    if issue_name in {MERGED_CONTAINER_ISSUE_NAME, "桶站及周边环境", "容器完好率", "桶站便利措施", "垃圾桶成组配备", "分类标识"}:
        return "投放点情况"
    if issue_name in {"垃圾分类驿站", "大件垃圾投放点设置", "装修垃圾投放点设置", "大件垃圾管理"}:
        return "大件垃圾投放点"
    if issue_name in {"公示牌", "责任公示"}:
        return "小区公示牌"
    if issue_name in {"宣传氛围"}:
        return "规范管理"
    return ""


def work_action_body(category: str, city_problems: Counter, district_problems: Counter) -> str:
    sentences: list[str] = []
    _rules, defaults, _warnings = load_work_action_config()
    for problem_counts in (city_problems, district_problems):
        for problem, count in problem_counts.most_common():
            if count <= 0:
                continue
            sentence = work_action_sentence(category, problem)
            if sentence and sentence not in sentences:
                sentences.append(sentence)
            if len(sentences) == 3:
                return "".join(sentences)
    return "".join(sentences) if sentences else defaults[category]


def work_action_sentence(category: str, problem: str) -> str:
    rules, defaults, _warnings = load_work_action_config()
    normalized_problem = clean_text(problem)
    for keywords, sentence in rules.get(category, []):
        if any(keyword in normalized_problem for keyword in keywords):
            return sentence
    return defaults[category]


def validate_template_variables(template_path: Path, context: dict[str, Any]) -> None:
    tpl = DocxTemplate(str(template_path))
    variables = tpl.get_undeclared_template_variables()
    missing = sorted(v for v in variables if v not in context)
    if missing:
        raise WeeklyReportError("模板变量缺失：" + "、".join(missing))


def write_issue_ledger(rows: list[dict[str, Any]], output_path: Path) -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "问题台账"
    headers = [
        "编号",
        "1级点位",
        "2级点位",
        "3级点位",
        "4级点位",
        "5级点位",
        "1级指标",
        "2级指标",
        "3级指标",
        "具体问题",
        "检查时间",
        "上报时间",
        "图片1",
        "图片2",
        "图片3",
        "图片4",
        "图片5",
    ]
    ws.append(headers)
    for idx, row in enumerate(rows, start=1):
        ws.append(
            [
                idx,
                clean_text(row.get("1级点位")),
                clean_text(row.get("2级点位")),
                clean_text(row.get("3级点位")),
                clean_text(row.get("4级点位")),
                clean_text(row.get("5级点位")),
                clean_text(row.get("1级指标")),
                clean_text(row.get("2级指标")),
                clean_text(row.get("3级指标")),
                clean_text(row.get("具体问题")),
                clean_text(row.get("检查时间")),
                clean_text(row.get("上报时间")),
                clean_text(row.get("图片1")),
                clean_text(row.get("图片2")),
                clean_text(row.get("图片3")),
                clean_text(row.get("图片4")),
                clean_text(row.get("图片5")),
            ]
        )
    wb.save(output_path)
    wb.close()


def render_report(template_path: Path, context: dict[str, Any], output_path: Path) -> None:
    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    tpl.save(output_path)
    merge_rank_group_cells(output_path)
    update_rank_chart(output_path, context["district_rank_rows"])


def merge_rank_group_cells(path: Path) -> None:
    doc = Document(str(path))
    if len(doc.tables) < 2:
        return
    table = doc.tables[1]
    if len(table.rows) <= 3:
        doc.save(str(path))
        return
    start = 2
    current = table.cell(start, 1).text.strip()
    for row_idx in range(3, len(table.rows) + 1):
        group = table.cell(row_idx, 1).text.strip() if row_idx < len(table.rows) else None
        if group != current:
            end = row_idx - 1
            if current and end > start:
                merged = table.cell(start, 1).merge(table.cell(end, 1))
                set_cell_text_style(merged, current, size_pt=11)
            start = row_idx
            current = group
    doc.save(str(path))


def set_cell_text_style(cell: Any, text: str, size_pt: int = 11) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.name = "仿宋_GB2312"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")


def update_rank_chart(path: Path, rank_rows: list[dict[str, Any]]) -> None:
    if not rank_rows:
        return
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        chart_names = [n for n in names if n.startswith("word/charts/chart") and n.endswith(".xml")]
        workbook_names = [n for n in names if n.startswith("word/embeddings/") and n.endswith(".xlsx")]
        if not chart_names or not workbook_names:
            return
        chart_name = chart_names[0]
        workbook_name = workbook_names[0]
        workbook_bytes = zin.read(workbook_name)
        chart_xml = zin.read(chart_name)

    updated_workbook = update_chart_workbook(workbook_bytes, rank_rows)
    updated_chart = update_chart_xml(chart_xml, rank_rows)

    tmp = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == workbook_name:
                data = updated_workbook
            elif item.filename == chart_name:
                data = updated_chart
            zout.writestr(item, data)
    tmp.replace(path)


def update_chart_workbook(workbook_bytes: bytes, rank_rows: list[dict[str, Any]]) -> bytes:
    ensure_project_temp()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx", dir=PROJECT_TMP) as src:
        src.write(workbook_bytes)
        src_path = Path(src.name)
    out_path = PROJECT_TMP / f"{src_path.stem}_updated.xlsx"
    try:
        wb = openpyxl.load_workbook(src_path)
        ws = wb.active
        for merged_range in list(ws.merged_cells.ranges):
            ws.unmerge_cells(str(merged_range))
        avg = sum(float(row["issue_rate"]) for row in rank_rows) / len(rank_rows)
        for row_idx in range(3, 64):
            for col_idx in range(1, 14):
                ws.cell(row_idx, col_idx).value = None
        for idx, row in enumerate(rank_rows, start=3):
            ws.cell(idx, 1).value = row["group"] if row.get("group_first") else None
            ws.cell(idx, 2).value = row["town"]
            ws.cell(idx, 3).value = float(row["issue_rate"])
            ws.cell(idx, 3).number_format = "0.00"
            ws.cell(idx, 4).value = avg
            ws.cell(idx, 4).number_format = "0.00"
            ws.cell(idx, 8).value = row["town"]
            ws.cell(idx, 9).value = row["issue_count"]
            ws.cell(idx, 10).value = row["issue_count"]
            ws.cell(idx, 11).value = row["checked_count"]
            ws.cell(idx, 12).value = float(row["issue_rate"])
            ws.cell(idx, 12).number_format = "0.00"
            ws.cell(idx, 13).value = f"=J{idx}/K{idx}"
            ws.cell(idx, 13).number_format = "0.00"
        total_row = 3 + len(rank_rows)
        ws.cell(total_row, 8).value = "总计"
        ws.cell(total_row, 9).value = f"=SUM(I3:I{total_row - 1})"
        ws.cell(total_row, 10).value = f"=SUM(J3:J{total_row - 1})"
        ws.cell(total_row, 11).value = f"=SUM(K3:K{total_row - 1})"
        ws.cell(total_row, 12).value = f"=J{total_row}/K{total_row}"
        ws.cell(total_row, 13).value = f"=J{total_row}/K{total_row}"
        wb.save(out_path)
        wb.close()
        return out_path.read_bytes()
    finally:
        for p in (src_path, out_path):
            try:
                p.unlink()
            except FileNotFoundError:
                pass


def update_chart_xml(chart_xml: bytes, rank_rows: list[dict[str, Any]]) -> bytes:
    ET.register_namespace("c", "http://schemas.openxmlformats.org/drawingml/2006/chart")
    ET.register_namespace("a", "http://schemas.openxmlformats.org/drawingml/2006/main")
    ET.register_namespace("r", "http://schemas.openxmlformats.org/officeDocument/2006/relationships")
    ET.register_namespace("mc", "http://schemas.openxmlformats.org/markup-compatibility/2006")
    ET.register_namespace("c14", "http://schemas.microsoft.com/office/drawing/2007/8/2/chart")
    ET.register_namespace("c15", "http://schemas.microsoft.com/office/drawing/2012/chart")
    ns = {"c": "http://schemas.openxmlformats.org/drawingml/2006/chart"}
    root = ET.fromstring(chart_xml)
    end_row = 2 + len(rank_rows)
    avg = sum(float(row["issue_rate"]) for row in rank_rows) / len(rank_rows)
    categories = [(row["group"] if row.get("group_first") else None, row["town"]) for row in rank_rows]
    rates = [float(row["issue_rate"]) for row in rank_rows]
    averages = [avg for _row in rank_rows]
    group_point_styles = chart_group_point_styles(root, ns)
    clear_chart_point_overrides(root, ns)
    apply_chart_group_colors(root, rank_rows, group_point_styles, ns)

    for cat in root.findall(".//c:cat", ns):
        ref = cat.find(".//c:multiLvlStrRef", ns)
        if ref is None:
            continue
        f = ref.find("c:f", ns)
        if f is not None:
            f.text = f"Sheet1!$A$3:$B${end_row}"
        cache = ref.find("c:multiLvlStrCache", ns)
        if cache is not None:
            replace_multi_level_cache(cache, categories, ns)

    val_nodes = root.findall(".//c:val", ns)
    if val_nodes:
        update_num_ref(val_nodes[0], f"Sheet1!$C$3:$C${end_row}", rates, ns)
    if len(val_nodes) > 1:
        update_num_ref(val_nodes[1], f"Sheet1!$D$3:$D${end_row}", averages, ns)
    update_value_axis_scale(root, rates, ns)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def chart_group_point_styles(root: ET.Element, ns: dict[str, str]) -> dict[str, ET.Element]:
    series = root.find(".//c:barChart/c:ser", ns)
    if series is None:
        return {}
    point_styles = []
    for point in series.findall("c:dPt", ns):
        style = point.find("c:spPr", ns)
        if style is not None:
            point_styles.append(style)
    if not point_styles:
        return {}
    return {
        "城区组": deepcopy(point_styles[0]),
        "山区组": deepcopy(point_styles[-1]),
    }


def apply_chart_group_colors(
    root: ET.Element,
    rank_rows: list[dict[str, Any]],
    group_styles: dict[str, ET.Element],
    ns: dict[str, str],
) -> None:
    series = root.find(".//c:barChart/c:ser", ns)
    if series is None or not group_styles:
        return
    insert_at = next(
        (idx for idx, child in enumerate(series) if child.tag == f"{{{ns['c']}}}dLbls"),
        len(series),
    )
    for point_index, row in enumerate(rank_rows):
        style = group_styles.get(clean_text(row.get("group")))
        if style is None:
            continue
        point = ET.Element(f"{{{ns['c']}}}dPt")
        idx = ET.SubElement(point, f"{{{ns['c']}}}idx")
        idx.set("val", str(point_index))
        invert = ET.SubElement(point, f"{{{ns['c']}}}invertIfNegative")
        invert.set("val", "0")
        point.append(deepcopy(style))
        series.insert(insert_at, point)
        insert_at += 1


def clear_chart_point_overrides(root: ET.Element, ns: dict[str, str]) -> None:
    for ser_index, ser in enumerate(root.findall(".//c:ser", ns)):
        for child in list(ser):
            if child.tag == f"{{{ns['c']}}}dPt":
                ser.remove(child)
        d_lbls = ser.find("c:dLbls", ns)
        if d_lbls is None:
            continue
        for child in list(d_lbls):
            if child.tag == f"{{{ns['c']}}}dLbl":
                d_lbls.remove(child)
        show_value = "1" if ser_index == 0 else "0"
        for tag, value in {
            "showLegendKey": "0",
            "showVal": show_value,
            "showCatName": "0",
            "showSerName": "0",
            "showPercent": "0",
            "showBubbleSize": "0",
        }.items():
            node = d_lbls.find(f"c:{tag}", ns)
            if node is None:
                node = ET.SubElement(d_lbls, f"{{{ns['c']}}}{tag}")
            node.set("val", value)


def update_value_axis_scale(root: ET.Element, values: list[float], ns: dict[str, str]) -> None:
    if not values:
        return
    axis_max = nice_axis_max(max(values))
    major_unit = nice_major_unit(axis_max)
    for val_ax in root.findall(".//c:valAx", ns):
        scaling = val_ax.find("c:scaling", ns)
        if scaling is None:
            scaling = ET.SubElement(val_ax, f"{{{ns['c']}}}scaling")
        max_node = scaling.find("c:max", ns)
        if max_node is None:
            max_node = ET.SubElement(scaling, f"{{{ns['c']}}}max")
        max_node.set("val", f"{axis_max:g}")
        min_node = scaling.find("c:min", ns)
        if min_node is None:
            min_node = ET.SubElement(scaling, f"{{{ns['c']}}}min")
        min_node.set("val", "0")
        major_node = val_ax.find("c:majorUnit", ns)
        if major_node is None:
            major_node = ET.SubElement(val_ax, f"{{{ns['c']}}}majorUnit")
        major_node.set("val", f"{major_unit:g}")


def nice_axis_max(value: float) -> float:
    if value <= 0:
        return 1
    padded = value * 1.25
    return nice_ceiling(padded)


def nice_major_unit(axis_max: float) -> float:
    return nice_ceiling(axis_max / 5)


def nice_ceiling(value: float) -> float:
    if value <= 0:
        return 1
    magnitude = 10 ** int(f"{value:e}".split("e")[1])
    normalized = value / magnitude
    for step in (1, 1.2, 1.5, 2, 2.5, 3, 4, 5, 6, 8, 10):
        if normalized <= step:
            result = step * magnitude
            return round(result, 6)
    return round(10 * magnitude, 6)


def replace_multi_level_cache(cache: ET.Element, categories: list[tuple[str | None, str]], ns: dict[str, str]) -> None:
    for child in list(cache):
        cache.remove(child)
    pt_count = ET.SubElement(cache, f"{{{ns['c']}}}ptCount")
    pt_count.set("val", str(len(categories)))
    town_level = ET.SubElement(cache, f"{{{ns['c']}}}lvl")
    for idx, (_group, town) in enumerate(categories):
        add_chart_pt(town_level, idx, town, ns)
    group_level = ET.SubElement(cache, f"{{{ns['c']}}}lvl")
    for idx, (group, _town) in enumerate(categories):
        if group:
            add_chart_pt(group_level, idx, group, ns)


def update_num_ref(val: ET.Element, formula: str, values: list[float], ns: dict[str, str]) -> None:
    num_ref = val.find("c:numRef", ns)
    if num_ref is None:
        return
    f = num_ref.find("c:f", ns)
    if f is not None:
        f.text = formula
    cache = num_ref.find("c:numCache", ns)
    if cache is None:
        cache = ET.SubElement(num_ref, f"{{{ns['c']}}}numCache")
    for child in list(cache):
        cache.remove(child)
    fmt = ET.SubElement(cache, f"{{{ns['c']}}}formatCode")
    fmt.text = "0.00_ "
    pt_count = ET.SubElement(cache, f"{{{ns['c']}}}ptCount")
    pt_count.set("val", str(len(values)))
    for idx, value in enumerate(values):
        add_chart_pt(cache, idx, f"{value:.12g}", ns)


def add_chart_pt(parent: ET.Element, idx: int, value: str, ns: dict[str, str]) -> None:
    pt = ET.SubElement(parent, f"{{{ns['c']}}}pt")
    pt.set("idx", str(idx))
    v = ET.SubElement(pt, f"{{{ns['c']}}}v")
    v.text = value


def generate_weekly_report(
    district_base_path: Path,
    city_ledger_path: Path,
    template_path: Path,
    group_path: Path | None,
    output_dir: Path,
    report_year: str = "2026",
    district_period: str = "",
    city_period: str = "",
    city_sheet: str | None = None,
    previous_district_issues: int | None = None,
) -> GenerationResult:
    ensure_project_temp()
    output_dir.mkdir(parents=True, exist_ok=True)
    district = read_district_base(district_base_path)
    city = read_city_ledger(city_ledger_path, city_sheet)
    groups, group_source = resolve_groups(group_path)
    district_period_short = parse_period_short(district_period) or district["period_short"]
    city_period_short = parse_period_short(city_period) or city["period_short"]
    if not district_period_short:
        raise WeeklyReportError("无法识别区级检查周期，请手动填写。")
    if not city_period_short:
        raise WeeklyReportError("无法识别市级检查周期，请手动填写。")
    district["rows"] = filter_rows_by_case_period(district["rows"], district_period_short)
    district["issue_rows"] = [row for row in district["rows"] if is_issue_row(row)]
    apply_city_period_filter(city, city_period_short)
    context = build_context(
        district=district,
        city=city,
        groups=groups,
        template_path=template_path,
        report_year=report_year,
        district_period_short=district_period_short,
        city_period_short=city_period_short,
        previous_district_issues=previous_district_issues,
    )
    report_path = output_dir / report_filename(report_year, district_period_short)
    temp_run_dir = Path(tempfile.mkdtemp(prefix="weekly_report_", dir=PROJECT_TMP))
    try:
        period_name = period_for_filename(district_period_short)
        issue_path = temp_run_dir / f"问题台账（{period_name}）.xlsx"
        summary_path = temp_run_dir / f"生成统计摘要（{period_name}）.json"
        write_issue_ledger(district["issue_rows"], issue_path)
        if len(district["issue_rows"]) != context_summary_issue_total(context):
            raise WeeklyReportError("生成的问题台账行数与区级问题总数不一致。")
        render_report(template_path, context, report_path)
        summary = {
            "report_path": str(report_path),
            "issue_ledger_path": str(issue_path),
            "city_sheet": city["sheet_name"],
            "group_source": group_source,
            "district_period_short": district_period_short,
            "city_period_short": city_period_short,
            "city_source_rows": city["source_row_count"],
            "city_filtered_rows": len(city["rows"]),
            "city_filtered_out_rows": city["filtered_out_count"],
            "district_total_issues": len(district["issue_rows"]),
            "previous_district_issues": previous_district_issues,
            "city_total_issues": sum(r["问题数量"] for r in city["rows"]),
            "district_rank_rows": context["district_rank_rows"],
            "city_six_categories": context["city_six_categories"],
            "next_work_items": context["next_work_items"],
            "warnings": context.get("warnings", []),
        }
        summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
        return GenerationResult(report_path, context)
    finally:
        shutil.rmtree(temp_run_dir, ignore_errors=True)


def context_summary_issue_total(context: dict[str, Any]) -> int:
    match = re.search(r"累计发现问题(\d+)处", context.get("district_issue_summary", ""))
    return int(match.group(1)) if match else -1


def validate_inputs(
    district_base_path: Path,
    city_ledger_path: Path,
    group_path: Path | None = None,
    city_sheet: str | None = None,
    district_period: str = "",
    city_period: str = "",
) -> dict[str, Any]:
    district = read_district_base(district_base_path)
    city = read_city_ledger(city_ledger_path, city_sheet)
    groups, group_source = resolve_groups(group_path)
    district_period_short = parse_period_short(district_period) or district["period_short"]
    city_period_short = parse_period_short(city_period) or city["period_short"]
    if not district_period_short:
        raise WeeklyReportError("无法识别区级检查周期，请手动填写。")
    if not city_period_short:
        raise WeeklyReportError("无法识别市级检查周期，请手动填写。")
    if district_period_short:
        district["rows"] = filter_rows_by_case_period(district["rows"], district_period_short)
        district["issue_rows"] = [row for row in district["rows"] if is_issue_row(row)]
    apply_city_period_filter(city, city_period_short)
    district_towns = {clean_text(r.get("3级点位")) for r in district["rows"] if clean_text(r.get("3级点位"))}
    missing_groups = sorted(t for t in district_towns if t not in groups)
    community_count, village_count, _town_counts = checked_point_counts(district["rows"])
    city_total = sum(r["问题数量"] for r in city["rows"])
    city_points = {(r["乡镇街道"], r["小区/村"]) for r in city["rows"] if r["乡镇街道"] and r["小区/村"]}
    rank_preview = district_rank_analysis(district["rows"], district["issue_rows"], groups)[:10] if not missing_groups else []
    city_preview, _paragraphs = city_six_category_analysis(city["rows"], city_total)
    return {
        "district_period_short": district_period_short,
        "city_period_short": city_period_short,
        "district_issue_total": len(district["issue_rows"]),
        "district_checked_total": community_count + village_count,
        "district_community_count": community_count,
        "district_village_count": village_count,
        "city_issue_total": city_total,
        "city_checked_points": len(city_points),
        "city_source_rows": city["source_row_count"],
        "city_filtered_rows": len(city["rows"]),
        "city_filtered_out_rows": city["filtered_out_count"],
        "city_sheet": city["sheet_name"],
        "group_source": group_source,
        "missing_groups": missing_groups,
        "rank_preview": rank_preview,
        "city_six_categories": city_preview,
    }
